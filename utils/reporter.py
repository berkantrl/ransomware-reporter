from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches
from dateutil.relativedelta import relativedelta
import seaborn as sns
from collections import defaultdict
from datetime import datetime
import json

LANG_TEXTS = {
    'tr': {
        'report_title': 'Aylık Ransomware Kurbanları Raporu',
        'sector_distribution': 'Sektörel Dağılım',
        'group_distribution': 'Grup Dağılımı',
        'country_distribution': 'Ülkelere Göre Dağılım',
        'continent_distribution': 'Kıtasal Dağılım',
        'sector_chart': 'Sektörel Dağılım Grafiği',
        'group_chart': 'Grup Dağılımı Grafiği',
        'country_chart': 'Ülkelere Göre Dağılım Grafiği',
        'continent_chart': 'Kıtasal Dağılım Grafiği',
        'case': 'vaka',
    },
    'en': {
        'report_title': 'Monthly Ransomware Victims Report',
        'sector_distribution': 'Sector Distribution',
        'group_distribution': 'Group Distribution',
        'country_distribution': 'Country Distribution',
        'continent_distribution': 'Continent Distribution',
        'sector_chart': 'Sector Distribution Chart',
        'group_chart': 'Group Distribution Chart',
        'country_chart': 'Country Distribution Chart',
        'continent_chart': 'Continent Distribution Chart',
        'case': 'case',
    }
}

def create_graph(data_dict, title, filename):
    labels = list(data_dict.keys())
    sizes = list(data_dict.values())

    plt.figure(figsize=(10, 6))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')
    plt.title(title)
    plt.savefig(filename)
    plt.close()

def create_bar_chart(data, title, xlabel, ylabel, file_name):
    plt.figure(figsize=(10, 6))
    sns.barplot(x=list(data.keys()), y=list(data.values()), palette='viridis')
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(file_name)
    plt.close()

def create_report(json_file_name, lang):
    previous_month = datetime.now() - relativedelta(months=1)
    
    sector_count = defaultdict(int)
    group_count = defaultdict(int)
    country_count = defaultdict(int)
    continent_count = defaultdict(int)

    try:
        with open(json_file_name, 'r', encoding='utf-8') as f:
            data = json.load(f)
            for entry in data:
                sector_count[entry['sector']] += 1
                group_count[entry['group_name']] += 1
                country_count[entry['country']] += 1
                continent_count[entry['continent']] += 1
    except FileNotFoundError:
        print("Data File Not Found!")
        return
    
    texts = LANG_TEXTS.get(lang, LANG_TEXTS['en'])
    doc = Document()
    doc.add_heading(texts['report_title'], 0)

    doc.add_heading(texts['sector_distribution'], level=1)
    for sector, count in sector_count.items():
        doc.add_paragraph(f'{sector}: {count} vaka')

    doc.add_heading(texts['group_distribution'], level=1)
    for group, count in group_count.items():
        doc.add_paragraph(f'{group}: {count} vaka')

    doc.add_heading(texts['country_distribution'], level=1)
    for country, count in country_count.items():
        doc.add_paragraph(f'{country}: {count} vaka')

    doc.add_heading(texts['continent_distribution'], level=1)
    for continent, count in continent_count.items():
        doc.add_paragraph(f'{continent}: {count} vaka')

    create_bar_chart(country_count, texts['country_chart'], 'Ülke', 'Vaka Sayısı', 'charts/country_distribution.png')
    create_bar_chart(group_count, texts['group_chart'], 'Grup', 'Kurban Sayısı', 'charts/group_distribution.png')
    create_bar_chart(sector_count, texts['sector_chart'], 'Sektör', 'Vaka Sayısı', 'charts/sector_distribution.png')
    create_graph(continent_count, texts['continent_chart'], 'charts/continent_distribution.png')


    doc.add_heading(texts['sector_chart'], level=2)
    doc.add_picture('charts/sector_distribution.png', width=Inches(5))

    doc.add_heading(texts['group_chart'], level=2)
    doc.add_picture('charts/group_distribution.png', width=Inches(5))

    doc.add_heading(texts['country_chart'], level=2)
    doc.add_picture('charts/country_distribution.png', width=Inches(5))

    doc.add_heading(texts['continent_chart'], level=2)
    doc.add_picture('charts/continent_distribution.png', width=Inches(5))

    filename = f'reports/Ransomware_report_{datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")}.docx'
    doc.save(filename)
    return filename
